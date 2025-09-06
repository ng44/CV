from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def create_cv_word(output_path, readme_path):
    # Create a new Word document
    doc = Document()

    # Set default font and style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read content from README.md
    with open(readme_path, "r") as f:
        lines = f.readlines()

    for line in lines:
        if line.strip():  # Skip empty lines
            if line.startswith("#"):
                heading = doc.add_heading(line.strip("# "), level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif line.startswith("**"):
                subheading = doc.add_heading(line.strip("** "), level=2)
                subheading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph = doc.add_paragraph(line.strip())
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the Word document
    doc.save(output_path)

# Generate the CV Word document
create_cv_word("Nikolaos_Georgiadis_CV.docx", "README.md")
